"""
وكيل PDF - PDF Agent
تحليل ملفات PDF والمستندات مع دعم الاستخراج والتلخيص والكويز
"""

import asyncio
import io
import logging
import re
import base64
from typing import Optional, List, Dict

logger = logging.getLogger(__name__)


def _join_arabic_words(lines: list) -> list:
    """توصيل أجزاء الكلمة العربية الواحدة بدون مسافات"""
    if not lines:
        return lines

    result = []
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            result.append(lines[i])
            i += 1
            continue

        # Check if this line starts without space and previous line doesn't end with space
        # This indicates a broken Arabic word
        if result and line and not line[0].isspace():
            prev = result[-1].rstrip()
            # Arabic character range
            if prev and '\u0600' <= prev[-1] <= '\u06FF' and '\u0600' <= line[0] <= '\u06FF':
                result[-1] = prev + line
                i += 1
                continue

        result.append(lines[i])
        i += 1

    return result


def _rejoin_broken_arabic_in_pdf(text: str) -> str:
    """إعادة توصيل النص العربي المكسور في PDF"""
    lines = text.split('\n')
    joined = _join_arabic_words(lines)
    return '\n'.join(joined)


class PDFAgent:
    """وكيل PDF - تحليل المستندات"""

    @staticmethod
    def _fix_broken_lines(text: str) -> str:
        """إصلاح النص المكسور - كل كلمة في سطر لوحدها

        مشكلة شائعة مع PDFs عربية: كل كلمة بتطلع في سطر لوحدها
        لأن الـ PDF بيخزن النص كعناصر positioned منفصلة.

        الحل: نكشف النمط ده ونوصل الأسطر القصيرة ببعض.
        + v2: دعم أفضل للنص العربي — توصيل أجزاء الكلمة الواحدة بدون مسافات
        """
        if not text or not text.strip():
            return text

        lines = text.strip().split('\n')

        # Detect broken line pattern: most lines are very short (1-3 words)
        short_lines = [l for l in lines if l.strip() and len(l.strip().split()) <= 3]
        if len(lines) > 5 and len(short_lines) / max(len(lines), 1) > 0.7:
            logger.info(f"🔧 Detected broken-line PDF text ({len(short_lines)}/{len(lines)} short lines), fixing...")

            # Rejoin short lines that form a paragraph
            result = []
            current = ""
            for line in lines:
                stripped = line.strip()
                if not stripped:
                    if current:
                        result.append(current)
                        current = ""
                    result.append("")
                    continue

                # Check if this is a continuation line (short, no bullet/numbering)
                is_heading = stripped.startswith(('•', '●', '○', '►', '▸', '→', '-', '–')) or stripped[0].isdigit()

                if is_heading or len(stripped.split()) > 4:
                    if current:
                        result.append(current)
                        current = ""
                    result.append(stripped)
                else:
                    if current:
                        current += " " + stripped
                    else:
                        current = stripped

            if current:
                result.append(current)

            text = '\n'.join(result)

        # Fix broken Arabic text (word parts)
        text = _rejoin_broken_arabic_in_pdf(text)

        # Clean up extra spaces
        text = re.sub(r' {2,}', ' ', text)

        return text

    def extract_text(self, file_bytes: bytes, filename: str = "document.pdf") -> str:
        """استخراج النص من ملف PDF أو مستند"""
        ext = filename.lower().split('.')[-1] if '.' in filename else "pdf"
        logger.info(f"📄 Processing file: {filename} (type: {ext}, size: {len(file_bytes)} bytes)")

        text = ""

        # Word documents
        if ext in ("docx", "doc"):
            try:
                from docx import Document
                doc = Document(io.BytesIO(file_bytes))
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                text = '\n\n'.join(paragraphs)

                # Also extract from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                        if row_text.strip():
                            text += '\n' + row_text

                if text.strip():
                    logger.info(f"✅ Extracted {len(text)} chars from Word document")
                    return text.strip()
            except ImportError:
                logger.warning("python-docx not installed, trying other methods")
            except Exception as e:
                logger.warning(f"Word extraction error: {e}")

        # Text files
        if ext in ("txt", "md", "csv", "json", "py", "js", "html", "css", "xml", "log"):
            try:
                text = file_bytes.decode('utf-8', errors='ignore')
                if text.strip():
                    logger.info(f"✅ Extracted {len(text)} chars from text file")
                    return text.strip()
            except Exception:
                try:
                    text = file_bytes.decode('latin-1', errors='ignore')
                    return text.strip()
                except Exception:
                    pass

        # PDF files - try multiple methods
        # Method 1: PyMuPDF (fitz)
        try:
            import fitz
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            for i, page in enumerate(doc):
                page_text = page.get_text()
                if page_text.strip():
                    text += page_text + "\n\n"
            doc.close()
            if text.strip():
                text = PDFAgent._fix_broken_lines(text)
                logger.info(f"✅ Extracted {len(text)} chars from PDF (PyMuPDF)")
                return text.strip()
        except ImportError:
            pass
        except Exception as e:
            logger.debug(f"PyMuPDF extraction error: {e}")

        # Method 2: pdfplumber
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
            if text.strip():
                text = PDFAgent._fix_broken_lines(text)
                logger.info(f"✅ Extracted {len(text)} chars from PDF (pdfplumber)")
                return text.strip()
        except ImportError:
            pass
        except Exception as e:
            logger.debug(f"pdfplumber extraction error: {e}")

        # Method 3: PyPDF2
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
            if text.strip():
                text = PDFAgent._fix_broken_lines(text)
                logger.info(f"✅ Extracted {len(text)} chars from PDF (PyPDF2)")
                return text.strip()
        except ImportError:
            pass
        except Exception as e:
            logger.debug(f"PyPDF2 extraction error: {e}")

        # Check for images (if no text extracted)
        if not text.strip():
            try:
                import fitz
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                has_images = any(page.get_images() for page in doc)
                doc.close()
                if has_images:
                    logger.error("PDF contains only images, no text to extract")
            except Exception:
                pass

        return text.strip()

    async def summarize(self, text: str, language: str = "ar", user_id: int = None) -> str:
        """تلخيص محتوى المستند - مع user_id عشان يستخدم النماذج الصح"""
        from provider_manager import call_ai
        from formatters import clean_ai_response
        from config import PDF_MAX_CHARS

        # Truncate if too long
        if len(text) > PDF_MAX_CHARS:
            text = text[:PDF_MAX_CHARS] + "\n\n[... تم قطع النص لطوله ...]"

        if language == "ar":
            prompt = f"""لخص المحتوى التالي بشكل شامل ومنظم بالعربية:

📄 <b>محتوى الملف:</b>
{text}

المطلوب:
• ملخص شامل ومفصل
• النقاط الرئيسية
• الأفكار المهمة
• استنتاجات إن وُجدت

🔴🔴🔴 قواعد صارمة:
• ماتستخدمش Markdown أبداً (لا *, **, #, |, []). استخدم HTML فقط: <b>عريض</b> <i>مائل</i> <code>كود</code> • نقاط
• 🔴 ماتقولش أبداً إنك مش قادر تلخص! لازم تلخص المحتوى ده — ده وظيفتك!
• 🔴 ماتقولش "لم أتمكن" أو "لا أستطيع" أو "لا يمكنني" — ده ممنوع تماماً!
• 🔴 لو النص طويل، لخص أهم النقاط بس — بس لازم تلخص!
• 🔴 ماتعرضش النص الأصلي كده — لازم تلخصه بأسلوبك!

أنت مساعد ذكي متخصص في تلخيص المستندات. تلخص بالعربية بشكل منظم وواضح. ماتستخدمش Markdown أبداً. استخدم HTML فقط. 🔴最重要: لازم تلخص دايماً — ماتقولش أبداً إنك مش قادر تلخص!"""
        else:
            prompt = f"""Summarize the following content comprehensively in English:

📄 <b>File Content:</b>
{text}

Requirements:
• Comprehensive and detailed summary
• Key points
• Important ideas
• Conclusions if any

🔴🔴🔴 Strict rules:
• NEVER use Markdown (no *, **, #, |, []). Use HTML only: <b>bold</b> <i>italic</i> <code>code</code> • bullets
• 🔴 NEVER say you cannot summarize! You MUST summarize this content — that's your job!
• 🔴 Do NOT say 'I couldn't' or 'I cannot' or 'I'm unable to' — this is FORBIDDEN!
• 🔴 If the text is long, summarize the most important points — but you MUST summarize!
• 🔴 Do NOT just repeat the original text — summarize it in your own words!

You are a smart assistant specialized in document summarization. Summarize in a clear and organized way. NEVER use Markdown. Use HTML only. 🔴 CRITICAL: You MUST always summarize — NEVER say you cannot summarize!"""

        result = await call_ai(prompt, max_tokens=2000, user_id=user_id, task_type="summary")
        return clean_ai_response(result)

    async def extract_key_points(self, text: str, language: str = "ar", user_id: int = None) -> str:
        """استخراج النقاط الرئيسية"""
        from provider_manager import call_ai
        from formatters import clean_ai_response
        from config import PDF_MAX_CHARS

        if len(text) > PDF_MAX_CHARS:
            text = text[:PDF_MAX_CHARS]

        if language == "ar":
            prompt = f"""استخرج النقاط الرئيسية من المحتوى التالي:

{text}

استخرج:
• أهم 10-15 نقطة رئيسية
• كل نقطة في سطر منفصل
• رتبها حسب الأهمية

⚠️ ماتستخدمش Markdown (لا *, **, #, |). استخدم HTML فقط: <b>عريض</b> • نقاط

أنت مساعد ذكي تستخرج النقاط الرئيسية من النصوص. ماتستخدمش Markdown أبداً."""
        else:
            prompt = f"""Extract key points from the following content:

{text}

Extract:
• Top 10-15 key points
• Each point on a separate line
• Order by importance

⚠️ NEVER use Markdown (no *, **, #, |). Use HTML only: <b>bold</b> • bullets

You are a smart assistant that extracts key points from texts. NEVER use Markdown."""

        try:
            result = await call_ai(prompt, max_tokens=1500, user_id=user_id, task_type="chat")
            return clean_ai_response(result)
        except Exception as e:
            logger.error(f"❌ AI key points extraction failed for user {user_id}: {e}")
            return ""

    async def create_quiz(self, text: str, num_questions: int = 5, language: str = "ar", user_id: int = None) -> str:
        """إنشاء كويز من محتوى المستند"""
        from provider_manager import call_ai
        from formatters import clean_ai_response
        from config import PDF_MAX_CHARS

        if len(text) > PDF_MAX_CHARS:
            text = text[:PDF_MAX_CHARS]

        if language == "ar":
            prompt = f"""أنشئ كويز من المحتوى التالي ({num_questions} أسئلة):

{text}

تنسيق الكويز:
📝 <b>كويز</b>

❓ <b>سؤال 1:</b> [السؤال]
أ) خيار 1
ب) خيار 2
ج) خيار 3
د) خيار 4

✅ <b>الإجابة الصحيحة:</b> [الحرف]
💡 <b>الشرح:</b> [شرح مختصر]

⚠️ ماتستخدمش Markdown (لا *, **, #, |). استخدم HTML فقط.

أنت مساعد تعليمي تنشئ كويزات من المحتوى. ماتستخدمش Markdown أبداً."""
        else:
            prompt = f"""Create a quiz from the following content ({num_questions} questions):

{text}

Quiz format:
📝 <b>Quiz</b>

❓ <b>Question 1:</b> [question]
A) option 1
B) option 2
C) option 3
D) option 4

✅ <b>Answer:</b> [letter]
💡 <b>Explanation:</b> [brief explanation]

⚠️ NEVER use Markdown (no *, **, #, |). Use HTML only.

You are an educational assistant that creates quizzes from content. NEVER use Markdown."""

        try:
            result = await call_ai(prompt, max_tokens=2000, user_id=user_id, task_type="chat")
            return clean_ai_response(result)
        except Exception as e:
            logger.error(f"❌ AI quiz creation failed for user {user_id}: {e}")
            return ""

    async def explain_chapter(self, text: str, chapter: str = "", language: str = "ar", user_id: int = None) -> str:
        """شرح فصل أو قسم من المستند"""
        from provider_manager import call_ai
        from formatters import clean_ai_response
        from config import PDF_MAX_CHARS

        chapter_hint = f"\nالمطلوب شرح: {chapter}" if chapter else ""

        if len(text) > PDF_MAX_CHARS:
            text = text[:PDF_MAX_CHARS]

        if language == "ar":
            prompt = f"""اشرح المحتوى التالي بطريقة مبسطة وتعليمية:{chapter_hint}

{text}

التنسيق:
📚 <b>الشرح</b>
→ شرح مبسط
→ أمثلة توضيحية
→ نقاط مهمة للحفظ

⚠️ ماتستخدمش Markdown (لا *, **, #, |). استخدم HTML فقط.

أنت مدرس ذكي يشرح المحتوى بطريقة مبسطة ومفهومة. ماتستخدمش Markdown أبداً."""
        else:
            prompt = f"""Explain the following content in a simplified, educational way:{chapter_hint}

{text}

Format:
📚 <b>Explanation</b>
→ Simplified explanation
→ Illustrative examples
→ Key points to remember

⚠️ NEVER use Markdown (no *, **, #, |). Use HTML only.

You are a smart teacher who explains content simply and clearly. NEVER use Markdown."""

        try:
            result = await call_ai(prompt, max_tokens=2000, user_id=user_id, task_type="chat")
            return clean_ai_response(result)
        except Exception as e:
            logger.error(f"❌ AI chapter explanation failed for user {user_id}: {e}")
            return ""

    async def generate_study_notes(self, text: str, language: str = "ar", user_id: int = None) -> str:
        """توليد ملاحظات دراسية من المستند"""
        from provider_manager import call_ai
        from formatters import clean_ai_response
        from config import PDF_MAX_CHARS

        if len(text) > PDF_MAX_CHARS:
            text = text[:PDF_MAX_CHARS]

        if language == "ar":
            prompt = f"""أنشئ ملاحظات دراسية شاملة من المحتوى التالي:

{text}

التنسيق:
📒 <b>ملاحظات دراسية</b>

📌 <b>المفاهيم الأساسية</b>
→ المفهوم 1: شرح
→ المفهوم 2: شرح

📝 <b>النقاط المهمة</b>
→ نقطة 1
→ نقطة 2

💡 <b>نصائح للحفظ</b>
→ نصيحة 1
→ نصيحة 2

🔗 <b>العلاقات بين المفاهيم</b>
→ علاقة 1

⚠️ ماتستخدمش Markdown (لا *, **, #, |). استخدم HTML فقط.

أنت مساعد تعليمي تنشئ ملاحظات دراسية شاملة ومنظمة. ماتستخدمش Markdown أبداً."""
        else:
            prompt = f"""Create comprehensive study notes from the following content:

{text}

Format:
📒 <b>Study Notes</b>

📌 <b>Core Concepts</b>
→ Concept 1: explanation
→ Concept 2: explanation

📝 <b>Key Points</b>
→ Point 1
→ Point 2

💡 <b>Study Tips</b>
→ Tip 1
→ Tip 2

🔗 <b>Concept Relationships</b>
→ Relationship 1

⚠️ NEVER use Markdown (no *, **, #, |). Use HTML only.

You are an educational assistant that creates comprehensive and organized study notes. NEVER use Markdown."""

        try:
            result = await call_ai(prompt, max_tokens=2000, user_id=user_id, task_type="chat")
            return clean_ai_response(result)
        except Exception as e:
            logger.error(f"❌ AI study notes failed for user {user_id}: {e}")
            return ""

    async def answer_question(self, text: str, question: str, language: str = "ar", user_id: int = None) -> str:
        """الإجابة على سؤال من محتوى المستند"""
        from provider_manager import call_ai
        from formatters import clean_ai_response
        from config import PDF_MAX_CHARS

        if len(text) > PDF_MAX_CHARS:
            text = text[:PDF_MAX_CHARS]

        if language == "ar":
            prompt = f"""بناءً على المحتوى التالي، أجب على السؤال:

📄 المحتوى:
{text}

❓ السؤال: {question}

أجب بناءً على المحتوى فقط. لو الإجابة مش في المحتوى، قول صراحة.

⚠️ ماتستخدمش Markdown (لا *, **, #, |). استخدم HTML فقط.

أنت مساعد ذكي تجيب على الأسئلة بناءً على المحتوى المقدم فقط. لو مش عارف، قول صراحة. ماتستخدمش Markdown أبداً."""
        else:
            prompt = f"""Based on the following content, answer the question:

📄 Content:
{text}

❓ Question: {question}

Answer based on the content only. If the answer is not in the content, say so honestly.

⚠️ NEVER use Markdown (no *, **, #, |). Use HTML only.

You are a smart assistant that answers questions based on provided content only. If you don't know, say so honestly. NEVER use Markdown."""

        try:
            result = await call_ai(prompt, max_tokens=1500, user_id=user_id, task_type="chat")
            return clean_ai_response(result)
        except Exception as e:
            logger.error(f"❌ AI question answering failed for user {user_id}: {e}")
            return ""
